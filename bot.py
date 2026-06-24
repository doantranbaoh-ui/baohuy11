# -*- coding: utf-8 -*-
# ┌────────────────────────────────────────────────────────────────────────┐
# │                    NÃO ROBOT - ULTIMATE EDITION                        │
# │  3000 dòng - Điểm danh + Nổ Hũ + Games + File + Voice + Stats + Admin │
# │  Tác giả: palofsc (palo)  |  Ngày: 2026-06-24                          │
# └────────────────────────────────────────────────────────────────────────┘
import sys, io, os, json, time, random, re, html, hashlib, subprocess
import socket, signal, logging, base64, tempfile, asyncio, traceback
import urllib.parse, urllib.request, zipfile, csv, xml.etree.ElementTree as ET
from threading import Thread, Lock, Event, Timer
from datetime import datetime, timedelta, date
from collections import deque, defaultdict, OrderedDict, Counter
from typing import Dict, List, Optional, Any, Tuple, Union, Callable
from concurrent.futures import ThreadPoolExecutor, as_completed
from queue import Queue, PriorityQueue
from dataclasses import dataclass, field
from io import StringIO, BytesIO

# ─── LOGGING ──────────────────────────────────────────────────────────────
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
                    handlers=[logging.StreamHandler()])
logger = logging.getLogger("NaoRobot")

# ─── ENCODING ──────────────────────────────────────────────────────────────
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding='utf-8')

# ─── KEEP ALIVE ────────────────────────────────────────────────────────────
try:
    from keep_alive import keep_alive
    keep_alive()
except ImportError:
    pass

# ─── THƯ VIỆN NGOÀI ───────────────────────────────────────────────────────
import telebot
from telebot import types, util
import requests
import pytz

# ─── THƯ VIỆN ĐỌC FILE ──────────────────────────────────────────────────
try:
    import PyPDF2
    HAS_PYPDF2 = True
except ImportError:
    HAS_PYPDF2 = False
    logger.warning("PyPDF2 chưa cài: pip install PyPDF2")

try:
    import docx
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    logger.warning("python-docx chưa cài: pip install python-docx")

try:
    from bs4 import BeautifulSoup
    HAS_BS4 = True
except ImportError:
    HAS_BS4 = False
    logger.warning("BeautifulSoup4 chưa cài: pip install beautifulsoup4")

try:
    import chardet
    HAS_CHARDET = True
except ImportError:
    HAS_CHARDET = False

# ╔══════════════════════════════════════════════════════════════╗
# ║  NÃO (BRAIN) - LỚP ĐIỀU KHIỂN TRUNG TÂM                    ║
# ╚══════════════════════════════════════════════════════════════╝
class Brain:
    def __init__(self, save_path: str = "brain.json"):
        self.save_path = save_path
        self.state: str = "normal"
        self.mood: int = 0
        self.learned: defaultdict = defaultdict(int)
        self.banned_words: set = set()
        self.trusted_users: set = set()
        self.stats: Dict[str, Any] = {
            "msg_processed": 0, "spam_blocked": 0, "ai_calls": 0,
            "errors": 0, "votes_created": 0, "voice_generated": 0,
            "files_processed": 0, "daily_checkins": 0, "nohu_spins": 0,
            "uptime_start": time.time(), "last_save": time.time()
        }
        self.decision_log: deque = deque(maxlen=200)
        self.last_health_check: float = time.time()
        self.repair_mode: bool = False
        self.file_lock = Lock()
        self.load_state()

    def load_state(self):
        if os.path.exists(self.save_path):
            try:
                with open(self.save_path, "r", encoding="utf-8") as f:
                    data = json.load(f)
                    self.learned = defaultdict(int, data.get("learned", {}))
                    self.banned_words = set(data.get("banned", []))
                    self.trusted_users = set(data.get("trusted", []))
                    saved = data.get("stats", {})
                    self.stats.update(saved)
                    self.stats["uptime_start"] = self.stats.get("uptime_start", time.time())
                    self.state = data.get("state", "normal")
                    self.mood = data.get("mood", 0)
            except: pass

    def save_state(self):
        with self.file_lock:
            self.stats["last_save"] = time.time()
            try:
                with open(self.save_path, "w", encoding="utf-8") as f:
                    json.dump({"learned": dict(self.learned), "banned": list(self.banned_words),
                               "trusted": list(self.trusted_users), "stats": self.stats,
                               "state": self.state, "mood": self.mood}, f, ensure_ascii=False, indent=2)
            except: self.stats["errors"] += 1

    def think(self, context: dict) -> str:
        uid = context.get("uid"); txt = context.get("txt", "")
        self.stats["msg_processed"] += 1
        words = re.findall(r'\b\w{3,}\b', txt.lower())
        for w in words: self.learned[w] += 1
        neg = ["bot ngu", "bot dở", "bot lỗi", "mày ngu"]
        pos = ["bot hay", "bot pro", "cảm ơn bot", "bot tốt"]
        if any(p in txt.lower() for p in neg): self.mood -= 2
        elif any(p in txt.lower() for p in pos): self.mood += 1
        self.mood = max(-10, min(10, self.mood))
        self.state = "aggressive" if self.mood < -5 else "normal"
        self.decision_log.append({"time": datetime.now(pytz.timezone('Asia/Ho_Chi_Minh')).strftime("%H:%M:%S"),
                                  "uid": uid, "decision": self.state, "mood": self.mood})
        if len(self.decision_log) % 5 == 0: self.save_state()
        return self.state

    def should_reply(self, uid: int, msg_text: str) -> bool:
        if uid in self.trusted_users: return True
        if self.learned.get(msg_text.lower(), 0) > 5: return random.random() > 0.3
        return random.random() > 0.1

    def get_insult_level(self) -> str:
        if self.state == "aggressive": return "extreme"
        elif self.mood < 0: return "high"
        return "normal"

    def health_check(self) -> str:
        now = time.time()
        if now - self.last_health_check > 300:
            self.last_health_check = now
            if self.stats["errors"] > 20:
                self.repair_mode = True; self.state = "repair"; self.stats["errors"] = 0
                return "repair"
            self.save_state()
        return "ok"

brain = Brain()

# ╔══════════════════════════════════════════════════════════════╗
# ║  CẤU HÌNH CHÍNH                                           ║
# ╚══════════════════════════════════════════════════════════════╝
TOKEN = os.getenv("BOT_TOKEN", "8080338995:AAEL2qb-TMjjUmoSvG1bWuY5M1QFST_zdJ4")
ADMIN_ID = int(os.getenv("ADMIN_ID", "5736655322"))
GROUP_ID = int(os.getenv("GROUP_ID", "-1003925717296"))

bot = telebot.TeleBot(TOKEN, num_threads=50)
tz = pytz.timezone('Asia/Ho_Chi_Minh')
ses = requests.Session()
ses.mount('https://', requests.adapters.HTTPAdapter(pool_connections=200, pool_maxsize=500, max_retries=3, pool_block=False))
ses.mount('http://', requests.adapters.HTTPAdapter(pool_connections=200, pool_maxsize=500, max_retries=3, pool_block=False))

ai_executor = ThreadPoolExecutor(max_workers=20, thread_name_prefix="AI")
voice_executor = ThreadPoolExecutor(max_workers=8, thread_name_prefix="Voice")
file_executor = ThreadPoolExecutor(max_workers=5, thread_name_prefix="File")
game_executor = ThreadPoolExecutor(max_workers=10, thread_name_prefix="Game")

# ╔══════════════════════════════════════════════════════════════╗
# ║  AI KEYS                                                   ║
# ╚══════════════════════════════════════════════════════════════╝
AI_KEYS = [
    {"key": "sk-d92be6f49626610cee386cf85897fe353cd5fadc44f66a73e98a0cce3efdfd8d", "url": "https://api.byesu.com/v1/chat/completions", "model": "gpt-4o", "status": True, "fail": 0, "last_used": 0},
    {"key": "sk-d1c9defa13eaa7386af8f711f38e9e8dd7a4754c9eebfe7f5642a391db82c2c3", "url": "https://api.byesu.com/v1/chat/completions", "model": "gpt-4o", "status": True, "fail": 0, "last_used": 0},
    {"key": "fe_oa_7bd49f79bc22bda1bc0c9b89f37741aa0a3086e87cfba034", "url": "https://api.freemodel.dev/v1/chat/completions", "model": "gpt-4o", "status": True, "fail": 0, "last_used": 0}
]
MAX_FAIL = 3; ck_idx = 0; ck_lock = Lock()

# ╔══════════════════════════════════════════════════════════════╗
# ║  KHO CHỬI                                                  ║
# ╚══════════════════════════════════════════════════════════════╝
KHO_NORMAL = ["Mồm thối, câm đi.", "Não bã đậu, im lặng.", "Thùng rỗng kêu to.", "Cào phím nhanh, não chậm.", "Ảo tưởng sức mạnh.", "Về nhà rửa bát.", "IQ âm, đừng nói.", "Không ai cần mày.", "Mày là gì? Không là gì.", "Câm mồm, đỡ nhục."]
KHO_HIGH = ["Nứt mắt đòi làm anh hùng.", "Đầu rỗng, mồm thối.", "Mạng xã hội nuôi mày à?", "Ra đời người ta vả cho.", "Mẹ gọi, về nhà đi.", "Tưởng mình ngầu? Hề vãi.", "Học không lo, cào phím giỏi.", "Tương lai mù mịt như chị Dậu.", "Đời vả mặt, mày cười ngây.", "Không có gì để nói với mày."]
KHO_EXTREME = ["Mày đáng giá bằng cái nút block.", "Tồn tại để làm gì? Để tao chửi à?", "Não mày như ổ đĩa format nhầm.", "Mày là lỗi của tự nhiên.", "Tao chửi mày còn thấy phí thời gian.", "Mày không đáng để tao nhớ tên.", "Cút về lỗ mà mày chui ra.", "Mày là minh chứng cho thất bại của tiến hóa.", "Tao nhìn mày mà tưởng đang xem phim hài.", "Mày sống làm gì?"]
def get_kho():
    lvl = brain.get_insult_level()
    if lvl == "extreme": return KHO_EXTREME
    elif lvl == "high": return KHO_HIGH
    return KHO_NORMAL

# ╔══════════════════════════════════════════════════════════════╗
# ║  BIẾN TOÀN CỤC & DỮ LIỆU                                   ║
# ╚══════════════════════════════════════════════════════════════╝
lock = Lock(); mem = deque(maxlen=50)
users: Dict[str, str] = {}
spam: Dict[int, List[float]] = {}
warn_counts: Dict[int, int] = {}
mutes: Dict[int, float] = {}
ai_cd: Dict[int, float] = {}
vote_active: Dict[int, Dict] = {}
game_sessions: Dict[int, Dict] = {}
file_cache: Dict[str, Dict] = {}

# ═══════════ HỆ THỐNG TIỀN TỆ ẢO ═══════════
# Tiền ảo trong group: Xu (💎)
user_balance: Dict[int, int] = {}          # uid -> số xu
daily_checkin: Dict[int, str] = {}         # uid -> ngày điểm danh cuối (isoformat)
BALANCE_FILE = "balances.json"
DAILY_FILE = "daily_checkins.json"

# ═══════════ NỔ HŨ ═══════════
nohu_jackpot: int = 0                      # Tổng tiền nổ hũ hiện tại
nohu_history: deque = deque(maxlen=20)     # Lịch sử nổ hũ
nohu_base = 100000                         # Jackpot khởi điểm
nohu_fee = 1000                            # Phí mỗi lần quay
nohu_multiplier = 0.05                     # 5% tiền quay vào jackpot
nohu_last_reset = time.time()
JACKPOT_FILE = "jackpot.json"

# ═══════════ THỐNG KÊ THÀNH VIÊN ═══════════
member_stats: Dict[str, Any] = {
    "daily_join": defaultdict(int), "daily_leave": defaultdict(int),
    "total_joined": 0, "total_left": 0, "current_members": 0,
    "join_dates": {}, "last_updated": time.time()
}

MAX_CACHE_SIZE = 50
MAX_FILE_SIZE = 20 * 1024 * 1024
USR_FILE = "usr.json"; STATS_FILE = "member_stats.json"; RULES_FILE = "rules.txt"
TELEGRAM_LINK = re.compile(r'(https?://)?(www\.)?(t\.me|telegram\.me|telegram\.org|tg\.me)/[a-zA-Z0-9_]{5,}|@[a-zA-Z0-9_]{5,}', re.I)

# ╔══════════════════════════════════════════════════════════════╗
# ║  TIỆN ÍCH CHUNG                                            ║
# ╚══════════════════════════════════════════════════════════════╝
def load_json(path: str, default: Any = {}) -> Any:
    if os.path.exists(path):
        try:
            with open(path, 'r', encoding='utf-8') as f: return json.load(f)
        except: pass
    return default

def save_json(path: str, data: Any) -> None:
    with lock:
        try:
            with open(path, 'w', encoding='utf-8') as f: json.dump(data, f, ensure_ascii=False, indent=2)
        except Exception as e: logger.error(f"Lỗi save {path}: {e}")

def load_users() -> Dict[str, str]: return load_json(USR_FILE, {})
def save_users(data: Dict[str, str]): save_json(USR_FILE, data)

def load_balances() -> Dict[int, int]:
    data = load_json(BALANCE_FILE, {})
    return {int(k): v for k, v in data.items()}

def save_balances(data: Dict[int, int]):
    save_json(BALANCE_FILE, {str(k): v for k, v in data.items()})

def load_daily_checkins() -> Dict[int, str]:
    data = load_json(DAILY_FILE, {})
    return {int(k): v for k, v in data.items()}

def save_daily_checkins(data: Dict[int, str]):
    save_json(DAILY_FILE, {str(k): v for k, v in data.items()})

def load_jackpot() -> int:
    data = load_json(JACKPOT_FILE, {"jackpot": nohu_base, "history": []})
    return data.get("jackpot", nohu_base)

def save_jackpot(jackpot: int):
    save_json(JACKPOT_FILE, {"jackpot": jackpot, "history": list(nohu_history), "last_reset": nohu_last_reset})

def load_member_stats() -> Dict:
    data = load_json(STATS_FILE, {"daily_join": {}, "daily_leave": {}, "total_joined": 0, "total_left": 0, "current_members": 0, "join_dates": {}, "last_updated": time.time()})
    data["daily_join"] = defaultdict(int, data.get("daily_join", {}))
    data["daily_leave"] = defaultdict(int, data.get("daily_leave", {}))
    return data

def save_member_stats():
    data = dict(member_stats)
    data["daily_join"] = dict(data["daily_join"])
    data["daily_leave"] = dict(data["daily_leave"])
    save_json(STATS_FILE, data)

def del_msg(chat_id: int, msg_id: int, delay: int = 60):
    Thread(target=lambda: (time.sleep(delay), bot.delete_message(chat_id, msg_id)), daemon=True).start()

def is_admin(chat_id: int, user_id: int) -> bool:
    try:
        for admin in bot.get_chat_administrators(chat_id):
            if admin.user.id == user_id: return True
    except: pass
    return False

def is_grp(m) -> bool: return m.chat.id == GROUP_ID

def extract_user_and_reason(message, bot_username: str) -> Tuple[Optional[int], str]:
    target = None; reason = ""
    if message.reply_to_message:
        target = message.reply_to_message.from_user.id
        parts = message.text.split(maxsplit=1)
        if len(parts) > 1: reason = parts[1]
    else:
        parts = message.text.split(maxsplit=1)
        if len(parts) > 1:
            arg = parts[1].strip()
            if arg.isdigit(): target = int(arg)
            else:
                m = re.match(r'@(\w+)', arg)
                if m:
                    try:
                        target = bot.get_chat_member(message.chat.id, m.group(0)).user.id
                        reason = arg[m.end():].strip()
                    except: pass
                else:
                    nm = re.search(r'\d+', arg)
                    if nm: target = int(nm.group()); reason = arg[nm.end():].strip()
    return target, reason

def parse_duration(reason: str) -> int:
    m = re.search(r'(\d+)\s*(h|m|s|p)', reason.lower())
    if m:
        num = int(m.group(1)); unit = m.group(2)
        if unit == 's': return num
        elif unit == 'm': return num * 60
        elif unit == 'h': return num * 3600
        elif unit == 'p': return num * 60
    return 3600

def get_user_balance(uid: int) -> int:
    """Lấy số dư của user, nếu chưa có thì tạo mới với 5000 xu."""
    if uid not in user_balance:
        user_balance[uid] = 5000
        save_balances(user_balance)
    return user_balance[uid]

def add_balance(uid: int, amount: int) -> int:
    """Thêm tiền cho user."""
    bal = get_user_balance(uid)
    user_balance[uid] = max(0, bal + amount)
    save_balances(user_balance)
    return user_balance[uid]

def deduct_balance(uid: int, amount: int) -> bool:
    """Trừ tiền, trả về True nếu đủ."""
    bal = get_user_balance(uid)
    if bal >= amount:
        user_balance[uid] = bal - amount
        save_balances(user_balance)
        return True
    return False

# ╔══════════════════════════════════════════════════════════════╗
# ║  HỆ THỐNG ĐIỂM DANH HÀNG NGÀY                              ║
# ╚══════════════════════════════════════════════════════════════╝
def get_daily_reward(uid: int, consecutive_days: int = 1) -> int:
    """Tính thưởng điểm danh dựa trên số ngày liên tiếp."""
    base_reward = 500
    bonus = min(consecutive_days - 1, 6) * 200  # Tối đa +1200 bonus
    return base_reward + bonus

@bot.message_handler(commands=['daily', 'diemdanh', 'checkin'])
def daily_checkin_cmd(m):
    """Điểm danh hàng ngày nhận xu."""
    if not is_grp(m) or antispam(m): return
    uid = m.from_user.id
    today = date.today().isoformat()
    users[str(uid)] = m.from_user.first_name; save_users(users)

    last_checkin = daily_checkin.get(uid, "")
    yesterday = (date.today() - timedelta(days=1)).isoformat()

    if last_checkin == today:
        bal = get_user_balance(uid)
        msg = bot.reply_to(m, f"❌ <b>{html.escape(m.from_user.first_name)}</b>, hôm nay bạn đã điểm danh rồi!\n💰 Số dư: <b>{bal:,}</b> xu\n⏰ Quay lại sau 0h để điểm danh tiếp.", parse_mode="HTML")
        del_msg(m.chat.id, msg.message_id, 15)
        return

    # Tính ngày liên tiếp
    consecutive = 1
    if last_checkin == yesterday:
        # Tìm số ngày liên tiếp trước đó (tối đa 7)
        d = date.today() - timedelta(days=1)
        for i in range(1, 7):
            check_date = (d - timedelta(days=i)).isoformat()
            if daily_checkin.get(uid) == check_date:
                consecutive += 1
            else:
                break
    else:
        consecutive = 1

    reward = get_daily_reward(uid, consecutive)
    daily_checkin[uid] = today
    save_daily_checkins(daily_checkin)
    add_balance(uid, reward)
    brain.stats["daily_checkins"] += 1

    # Streak emoji
    streak_emoji = ["", "🔥", "🔥🔥", "💥", "💥💥", "⚡", "👑"]
    streak_str = streak_emoji[min(consecutive-1, 6)] if consecutive > 1 else ""

    response = (
        f"✅ <b>ĐIỂM DANH THÀNH CÔNG!</b>\n"
        f"👤 <b>{html.escape(m.from_user.first_name)}</b>\n"
        f"💰 Nhận: <b>+{reward:,}</b> xu\n"
        f"📅 Chuỗi: <b>{consecutive}</b> ngày liên tiếp {streak_str}\n"
        f"💎 Số dư: <b>{get_user_balance(uid):,}</b> xu\n"
        f"<i>Điểm danh mỗi ngày để nhận thưởng lớn hơn!</i>"
    )
    msg = bot.reply_to(m, response, parse_mode="HTML")
    del_msg(m.chat.id, msg.message_id, 30)

@bot.message_handler(commands=['balance', 'xu', 'money'])
def balance_cmd(m):
    """Xem số dư xu."""
    if not is_grp(m): return
    uid = m.from_user.id
    users[str(uid)] = m.from_user.first_name; save_users(users)
    
    # Nếu mention người khác
    target = uid
    target_name = m.from_user.first_name
    if m.reply_to_message:
        target = m.reply_to_message.from_user.id
        target_name = m.reply_to_message.from_user.first_name
    
    bal = get_user_balance(target)
    msg = bot.reply_to(m, f"💎 <b>{html.escape(target_name)}</b> có <b>{bal:,}</b> xu.", parse_mode="HTML")
    del_msg(m.chat.id, msg.message_id, 20)

@bot.message_handler(commands=['top', 'bxh'])
def top_balance_cmd(m):
    """Bảng xếp hạng người giàu nhất."""
    if not is_grp(m): return
    sorted_balances = sorted(user_balance.items(), key=lambda x: x[1], reverse=True)[:10]
    text = "🏆 <b>BẢNG XẾP HẠNG GIÀU NHẤT</b>\n━━━━━━━━━━━━━━━━━━━━\n"
    medals = ["🥇", "🥈", "🥉"] + ["  "] * 7
    for i, (uid, bal) in enumerate(sorted_balances):
        name = users.get(str(uid), str(uid))
        text += f"{medals[i]} <b>#{i+1}</b> <a href='tg://user?id={uid}'>{html.escape(name)}</a>: <code>{bal:,}</code> xu\n"
    if not sorted_balances:
        text += "Chưa có ai trong bảng xếp hạng."
    msg = bot.reply_to(m, text, parse_mode="HTML")
    del_msg(m.chat.id, msg.message_id, 45)

@bot.message_handler(commands=['give', 'chuyen'])
def give_money_cmd(m):
    """Chuyển xu cho người khác. /give [@mention/reply] [số xu]"""
    if not is_grp(m) or antispam(m): return
    uid = m.from_user.id
    users[str(uid)] = m.from_user.first_name; save_users(users)
    
    target = None; amount = 0
    if m.reply_to_message:
        target = m.reply_to_message.from_user.id
        parts = m.text.split()
        if len(parts) >= 2:
            try: amount = int(parts[1])
            except: amount = 0
    else:
        parts = m.text.split()
        if len(parts) >= 3:
            mention_match = re.match(r'@(\w+)', parts[1])
            if mention_match:
                try:
                    target = bot.get_chat_member(m.chat.id, parts[1]).user.id
                except: target = None
            elif parts[1].isdigit():
                target = int(parts[1])
            try: amount = int(parts[2])
            except: amount = 0
    
    if not target or target == uid:
        msg = bot.reply_to(m, "❌ Cần reply/@mention người nhận + số xu.", parse_mode="HTML"); del_msg(m.chat.id, msg.message_id, 10); return
    if amount < 100:
        msg = bot.reply_to(m, "❌ Số xu tối thiểu là 100.", parse_mode="HTML"); del_msg(m.chat.id, msg.message_id, 10); return
    
    # Phí chuyển 5%
    fee = int(amount * 0.05)
    total = amount + fee
    
    if not deduct_balance(uid, total):
        bal = get_user_balance(uid)
        msg = bot.reply_to(m, f"❌ Không đủ xu! Cần <b>{total:,}</b> (gồm {fee:,} phí). Số dư: <b>{bal:,}</b>", parse_mode="HTML"); del_msg(m.chat.id, msg.message_id, 15); return
    
    add_balance(target, amount)
    target_name = users.get(str(target), str(target))
    response = (
        f"💸 <b>CHUYỂN TIỀN</b>\n"
        f"👤 {html.escape(m.from_user.first_name)} → <a href='tg://user?id={target}'>{html.escape(target_name)}</a>\n"
        f"💰 Số tiền: <b>{amount:,}</b> xu\n"
        f"📊 Phí: <b>{fee:,}</b> xu\n"
        f"💎 Số dư còn: <b>{get_user_balance(uid):,}</b> xu"
    )
    msg = bot.reply_to(m, response, parse_mode="HTML")
    del_msg(m.chat.id, msg.message_id, 30)

# ╔══════════════════════════════════════════════════════════════╗
# ║  NỔ HŨ (JACKPOT SLOT MACHINE)                              ║
# ╚══════════════════════════════════════════════════════════════╝
# Biểu tượng nổ hũ
SLOT_SYMBOLS = ["🍒", "🍋", "🍊", "🍇", "💎", "🔔", "7️⃣"]
SLOT_WEIGHTS = [30, 25, 20, 15, 5, 3, 2]  # Tỉ lệ xuất hiện
SLOT_PAYOUTS = {
    "🍒🍒🍒": 5,     # x5 tiền cược
    "🍋🍋🍋": 8,
    "🍊🍊🍊": 12,
    "🍇🍇🍇": 20,
    "💎💎💎": 50,
    "🔔🔔🔔": 100,
    "7️⃣7️⃣7️⃣": 500,  # JACKPOT!
}

@bot.message_handler(commands=['nohu', 'slot', 'quay'])
def nohu_cmd(m):
    """Nổ Hũ: /nohu [cược] - Quay slot, cơ hội trúng JACKPOT!"""
    if not is_grp(m) or antispam(m): return
    uid = m.from_user.id
    users[str(uid)] = m.from_user.first_name; save_users(users)
    
    parts = m.text.split()
    if len(parts) < 2:
        jackpot = load_jackpot()
        msg = bot.reply_to(m, f"🎰 <b>NỔ HŨ</b>\n💰 <b>JACKPOT: {jackpot:,} xu</b>\n🎮 Dùng: /nohu [cược]\n💎 Phí: {nohu_fee:,} xu/lần\n🏆 3x7️⃣ = JACKPOT!", parse_mode="HTML")
        del_msg(m.chat.id, msg.message_id, 20)
        return
    
    try:
        bet = int(parts[1])
    except:
        msg = bot.reply_to(m, "❌ Cược phải là số.", parse_mode="HTML"); del_msg(m.chat.id, msg.message_id, 5); return
    
    if bet < 100:
        msg = bot.reply_to(m, "❌ Cược tối thiểu 100 xu.", parse_mode="HTML"); del_msg(m.chat.id, msg.message_id, 5); return
    if bet > 100000:
        msg = bot.reply_to(m, "❌ Cược tối đa 100,000 xu.", parse_mode="HTML"); del_msg(m.chat.id, msg.message_id, 5); return
    
    total_cost = bet + nohu_fee
    if not deduct_balance(uid, total_cost):
        bal = get_user_balance(uid)
        msg = bot.reply_to(m, f"❌ Không đủ xu! Cần <b>{total_cost:,}</b> (cược {bet:,} + phí {nohu_fee:,}). Số dư: <b>{bal:,}</b>", parse_mode="HTML"); del_msg(m.chat.id, msg.message_id, 10); return
    
    # Cập nhật jackpot (5% tiền cược vào quỹ)
    jackpot = load_jackpot()
    jackpot_contribution = int(bet * nohu_multiplier)
    jackpot += jackpot_contribution
    save_jackpot(jackpot)
    brain.stats["nohu_spins"] += 1
    
    # Quay 3 cột
    col1 = random.choices(SLOT_SYMBOLS, weights=SLOT_WEIGHTS, k=1)[0]
    col2 = random.choices(SLOT_SYMBOLS, weights=SLOT_WEIGHTS, k=1)[0]
    col3 = random.choices(SLOT_SYMBOLS, weights=SLOT_WEIGHTS, k=1)[0]
    result = f"{col1}{col2}{col3}"
    
    # Kiểm tra kết quả
    if col1 == col2 == col3:
        symbol = col1
        if symbol == "7️⃣":
            # JACKPOT!!!
            win_amount = jackpot
            add_balance(uid, win_amount)
            nohu_history.append({"uid": uid, "name": m.from_user.first_name, "amount": win_amount, "time": datetime.now(tz).strftime("%H:%M %d/%m")})
            save_jackpot(nohu_base)  # Reset jackpot
            outcome = f"🎉🎉🎉 <b>JACKPOT!!!</b> +{win_amount:,} xu"
            emoji = "🏆"
        else:
            multiplier = SLOT_PAYOUTS.get(result, 2)
            win_amount = bet * multiplier
            add_balance(uid, win_amount)
            outcome = f"✅ <b>NỔ HŨ!</b> (x{multiplier}) +{win_amount:,} xu"
            emoji = "🎉"
    elif col1 == col2 or col2 == col3 or col1 == col3:
        # 2 biểu tượng giống: hoàn 50% tiền cược
        win_amount = int(bet * 0.5)
        add_balance(uid, win_amount)
        outcome = f"🔹 2 giống: hoàn {win_amount:,} xu"
        emoji = "🔄"
    else:
        win_amount = 0
        outcome = f"💀 Thua -{total_cost:,} xu"
        emoji = "❌"
    
    response = (
        f"{emoji} <b>NỔ HŨ</b>\n"
        f"┌──────────┐\n"
        f"│ {col1}  {col2}  {col3} │\n"
        f"└──────────┘\n"
        f"🎯 Kết quả: {outcome}\n"
        f"💰 JACKPOT hiện tại: <b>{load_jackpot():,}</b> xu\n"
        f"💎 Số dư: <b>{get_user_balance(uid):,}</b> xu"
    )
    msg = bot.reply_to(m, response, parse_mode="HTML")
    del_msg(m.chat.id, msg.message_id, 30)

@bot.message_handler(commands=['jackpot', 'jp'])
def jackpot_cmd(m):
    """Xem jackpot hiện tại và lịch sử nổ hũ."""
    if not is_grp(m): return
    jackpot = load_jackpot()
    text = f"🎰 <b>NỔ HŨ JACKPOT</b>\n💰 <b>Tổng: {jackpot:,} xu</b>\n🎮 /nohu [cược] để quay!\n\n📜 <b>LỊCH SỬ NỔ HŨ:</b>\n"
    if nohu_history:
        for h in list(nohu_history)[-5:]:
            text += f"  🏆 {h['name']} trúng <b>{h['amount']:,}</b> xu ({h['time']})\n"
    else:
        text += "  Chưa có ai trúng JACKPOT.\n"
    text += f"\n<pre>3x7️⃣ = JACKPOT | 3x💎 = x50 | 3x🔔 = x100</pre>"
    msg = bot.reply_to(m, text, parse_mode="HTML")
    del_msg(m.chat.id, msg.message_id, 30)

# ╔══════════════════════════════════════════════════════════════╗
# ║  GAMES (TÀI XỈU, BẦU CUA, KÉO BÚA BAO, ĐOÁN SỐ)          ║
# ╚══════════════════════════════════════════════════════════════╝
GAME_STATES: Dict[int, Dict] = {}

def init_game_state(uid: int, game_type: str) -> Dict:
    if game_type == "taixiu":
        game = {"type": "taixiu", "balance": 1000, "started_at": time.time(), "history": [], "wins": 0, "losses": 0}
    elif game_type == "baucua":
        game = {"type": "baucua", "balance": 1000, "started_at": time.time(), "symbols": ["🦀", "🐟", "🦐", "🐓", "🦌", "🎃"], "history": [], "wins": 0, "losses": 0}
    elif game_type == "keobuabao":
        game = {"type": "keobuabao", "score": 0, "bot_score": 0, "draws": 0, "history": [], "started_at": time.time()}
    elif game_type == "doanso":
        game = {"type": "doanso", "secret": random.randint(1, 100), "attempts": 0, "max_attempts": 7, "started_at": time.time(), "guesses": []}
    else:
        return {}
    GAME_STATES[uid] = game
    return game

@bot.message_handler(commands=['taixiu'])
def taixiu_cmd(m):
    if not is_grp(m) or antispam(m): return
    uid = m.from_user.id
    users[str(uid)] = m.from_user.first_name; save_users(users)
    parts = m.text.split()
    
    if len(parts) < 3:
        if uid not in GAME_STATES or GAME_STATES[uid].get("type") != "taixiu": init_game_state(uid, "taixiu")
        g = GAME_STATES.get(uid, {})
        msg = bot.reply_to(m, f"🎲 <b>TÀI XỈU</b>\n/taixiu [tai/xiu] [cược]\n💎 Số dư game: <b>{g.get('balance', 1000)}</b> xu\nTài (11-18) | Xỉu (3-10)", parse_mode="HTML")
        del_msg(m.chat.id, msg.message_id, 20); return
    
    choice = parts[1].lower()
    try: bet = int(parts[2])
    except: bot.reply_to(m, "❌ Cược phải là số.", parse_mode="HTML"); return
    
    if choice not in ['tai', 'xiu']: bot.reply_to(m, "❌ Chọn 'tai' hoặc 'xiu'.", parse_mode="HTML"); return
    if uid not in GAME_STATES or GAME_STATES[uid].get("type") != "taixiu": init_game_state(uid, "taixiu")
    g = GAME_STATES[uid]
    if bet > g["balance"] or bet < 1: bot.reply_to(m, f"❌ Số dư game không đủ ({g['balance']} xu). Nạp thêm bằng /daily.", parse_mode="HTML"); return
    
    dice = [random.randint(1, 6) for _ in range(3)]
    total = sum(dice)
    result = "tai" if total >= 11 else "xiu"
    dice_str = " ".join([["⚀","⚁","⚂","⚃","⚄","⚅"][d-1] for d in dice])
    
    if choice == result:
        g["balance"] += bet; g["wins"] += 1
        outcome = f"✅ THẮNG +{bet} xu"
    else:
        g["balance"] -= bet; g["losses"] += 1
        outcome = f"❌ THUA -{bet} xu"
    
    g["history"].append({"dice": dice, "total": total, "result": result, "choice": choice, "bet": bet})
    if len(g["history"]) > 10: g["history"].pop(0)
    
    response = f"🎲 <b>TÀI XỈU</b>\n🎲 Xúc xắc: {dice_str} = <b>{total}</b> → <b>{result.upper()}</b>\n🎯 Bạn chọn: <b>{choice.upper()}</b>\n💰 {outcome} | Số dư game: <b>{g['balance']}</b> xu\n📊 Thắng: {g['wins']} | Thua: {g['losses']}"
    msg = bot.reply_to(m, response, parse_mode="HTML"); del_msg(m.chat.id, msg.message_id, 30)

@bot.message_handler(commands=['baucua'])
def baucua_cmd(m):
    if not is_grp(m) or antispam(m): return
    uid = m.from_user.id
    users[str(uid)] = m.from_user.first_name; save_users(users)
    parts = m.text.split()
    symbol_map = {"bau": 0, "bầu": 0, "cua": 1, "ca": 2, "cá": 2, "tom": 3, "tôm": 3, "ga": 4, "gà": 4, "nai": 5, "huou": 5, "hươu": 5}
    game_symbols = ["🦀 Bầu", "🐟 Cua", "🦐 Cá", "🐓 Tôm", "🦌 Gà", "🎃 Nai"]
    
    if len(parts) < 3:
        if uid not in GAME_STATES or GAME_STATES[uid].get("type") != "baucua": init_game_state(uid, "baucua")
        g = GAME_STATES.get(uid, {})
        msg = bot.reply_to(m, f"🎲 <b>BẦU CUA</b>\n{' | '.join(game_symbols)}\n/baucua [con] [cược]\n💎 Số dư game: <b>{g.get('balance', 1000)}</b> xu", parse_mode="HTML")
        del_msg(m.chat.id, msg.message_id, 25); return
    
    choice = parts[1].lower()
    try: bet = int(parts[2])
    except: bot.reply_to(m, "❌ Cược phải là số.", parse_mode="HTML"); return
    
    if choice not in symbol_map: bot.reply_to(m, f"❌ Chọn: {', '.join(symbol_map.keys())}", parse_mode="HTML"); return
    if uid not in GAME_STATES or GAME_STATES[uid].get("type") != "baucua": init_game_state(uid, "baucua")
    g = GAME_STATES[uid]
    if bet > g["balance"] or bet < 1: bot.reply_to(m, f"❌ Số dư game không đủ ({g['balance']} xu).", parse_mode="HTML"); return
    
    choice_idx = symbol_map[choice]
    roll = [random.randint(0, 5) for _ in range(3)]
    roll_symbols = [g["symbols"][i] for i in roll]
    matches = roll.count(choice_idx)
    
    if matches > 0:
        win_amount = bet * (matches + 1)
        g["balance"] += win_amount - bet; g["wins"] += 1
        outcome = f"✅ THẮNG +{win_amount - bet} xu (trúng {matches} con)"
    else:
        g["balance"] -= bet; g["losses"] += 1
        outcome = f"❌ THUA -{bet} xu"
    
    response = f"🎲 <b>BẦU CUA</b>\n🎯 Kết quả: {' '.join(roll_symbols)}\n🎯 Bạn chọn: <b>{g['symbols'][choice_idx]}</b> (trúng {matches}/3)\n💰 {outcome} | Số dư game: <b>{g['balance']}</b> xu"
    msg = bot.reply_to(m, response, parse_mode="HTML"); del_msg(m.chat.id, msg.message_id, 30)

@bot.message_handler(commands=['kbb', 'keobuabao'])
def kbb_cmd(m):
    if not is_grp(m) or antispam(m): return
    uid = m.from_user.id
    users[str(uid)] = m.from_user.first_name; save_users(users)
    parts = m.text.split()
    choices = {"keo": "✌️ Kéo", "kéo": "✌️ Kéo", "bua": "🔨 Búa", "búa": "🔨 Búa", "bao": "📄 Bao"}
    
    if len(parts) < 2:
        if uid not in GAME_STATES or GAME_STATES[uid].get("type") != "keobuabao": init_game_state(uid, "keobuabao")
        g = GAME_STATES.get(uid, {})
        msg = bot.reply_to(m, f"✌️ <b>KÉO BÚA BAO</b>\n/kbb [keo/bua/bao]\n👤 Bạn: <b>{g.get('score', 0)}</b> | 🤖 Bot: <b>{g.get('bot_score', 0)}</b> | 🤝 Hòa: <b>{g.get('draws', 0)}</b>", parse_mode="HTML")
        del_msg(m.chat.id, msg.message_id, 15); return
    
    choice = parts[1].lower()
    if choice not in choices: bot.reply_to(m, "❌ Chọn: keo/bua/bao", parse_mode="HTML"); return
    if uid not in GAME_STATES or GAME_STATES[uid].get("type") != "keobuabao": init_game_state(uid, "keobuabao")
    g = GAME_STATES[uid]
    
    user_choice = choices[choice]
    bot_choice = random.choice(list(choices.values()))
    user_idx = list(choices.values()).index(user_choice)
    bot_idx = list(choices.values()).index(bot_choice)
    
    if user_idx == bot_idx: result = "🤝 HÒA"; g["draws"] += 1
    elif (user_idx == 0 and bot_idx == 2) or (user_idx == 1 and bot_idx == 0) or (user_idx == 2 and bot_idx == 1): result = "✅ THẮNG"; g["score"] += 1
    else: result = "❌ THUA"; g["bot_score"] += 1
    
    response = f"✌️ <b>KÉO BÚA BAO</b>\n👤 Bạn: {user_choice}\n🤖 Bot: {bot_choice}\n📊 {result}\n🏆 Bạn: <b>{g['score']}</b> | Bot: <b>{g['bot_score']}</b> | Hòa: <b>{g['draws']}</b>"
    msg = bot.reply_to(m, response, parse_mode="HTML"); del_msg(m.chat.id, msg.message_id, 25)

@bot.message_handler(commands=['doanso'])
def doanso_cmd(m):
    if not is_grp(m) or antispam(m): return
    uid = m.from_user.id
    users[str(uid)] = m.from_user.first_name; save_users(users)
    parts = m.text.split()
    
    if len(parts) < 2:
        init_game_state(uid, "doanso")
        msg = bot.reply_to(m, "🔢 <b>ĐOÁN SỐ</b> (1-100)\n/doanso [số]\nBạn có <b>7</b> lần đoán!", parse_mode="HTML")
        del_msg(m.chat.id, msg.message_id, 15); return
    
    try: guess = int(parts[1])
    except: bot.reply_to(m, "❌ Nhập số từ 1-100.", parse_mode="HTML"); return
    if guess < 1 or guess > 100: bot.reply_to(m, "❌ Số từ 1-100.", parse_mode="HTML"); return
    if uid not in GAME_STATES or GAME_STATES[uid].get("type") != "doanso": init_game_state(uid, "doanso")
    g = GAME_STATES[uid]
    g["attempts"] += 1; g["guesses"].append(guess); secret = g["secret"]
    
    if guess == secret:
        reward = (8 - g["attempts"]) * 500
        add_balance(uid, reward)
        response = f"🎉 <b>CHÍNH XÁC!</b> Số là <b>{secret}</b> sau {g['attempts']} lần!\n💰 Thưởng: <b>+{reward:,}</b> xu\n🏆 {['⭐','🌟🌟','🌟🌟🌟','🏅','🏅🏅','👑','💎'][min(g['attempts']-1,6)]}"
        del GAME_STATES[uid]
    elif g["attempts"] >= g["max_attempts"]:
        response = f"💀 <b>HẾT LƯỢT!</b> Số là <b>{secret}</b>. Đã đoán: {', '.join(map(str, g['guesses']))}"
        del GAME_STATES[uid]
    elif guess < secret: response = f"🔢 <b>{guess}</b> → ⬆️ CAO HƠN ({g['max_attempts'] - g['attempts']} lần còn lại)"
    else: response = f"🔢 <b>{guess}</b> → ⬇️ THẤP HƠN ({g['max_attempts'] - g['attempts']} lần còn lại)"
    
    msg = bot.reply_to(m, response, parse_mode="HTML"); del_msg(m.chat.id, msg.message_id, 30)

# ╔══════════════════════════════════════════════════════════════╗
# ║  GOOGLE TTS VOICE                                          ║
# ╚══════════════════════════════════════════════════════════════╝
GOOGLE_TTS_URL = "https://translate.google.com/translate_tts"
GOOGLE_TTS_HEADERS = {"User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36", "Accept": "audio/mpeg, audio/*;q=0.9", "Referer": "https://translate.google.com/"}
MAX_CHUNK_SIZE = 180

@dataclass
class VoiceRequest:
    chat_id: int; reply_id: int; text: str; user_name: str; lang: str = "vi"
    created_at: float = field(default_factory=time.time)

voice_queue: Queue = Queue(maxsize=50)

def fetch_google_tts_chunk(text: str, lang: str = "vi") -> Optional[bytes]:
    params = {"ie": "UTF-8", "q": text, "tl": lang, "total": "1", "idx": "0", "textlen": str(len(text)), "client": "tw-ob", "prev": "input", "ttsspeed": "1.0"}
    try:
        resp = ses.get(GOOGLE_TTS_URL, params=params, headers=GOOGLE_TTS_HEADERS, timeout=15)
        if resp.status_code == 200 and len(resp.content) > 100: return resp.content
    except: pass
    return None

def split_text_into_chunks(text: str, max_size: int = MAX_CHUNK_SIZE) -> List[str]:
    if len(text) <= max_size: return [text]
    chunks = []; separators = ['. ', '! ', '? ', ', ', '; ', ': ', ' - ', '\n', ' ']
    while len(text) > max_size:
        best_pos = max_size
        for sep in separators:
            pos = text.rfind(sep, 0, max_size)
            if pos > max_size // 2: best_pos = pos + len(sep); break
        if best_pos > max_size or best_pos <= max_size // 3: best_pos = max_size
        chunks.append(text[:best_pos].strip()); text = text[best_pos:].strip()
    if text: chunks.append(text)
    return chunks

def generate_voice_google(text: str, lang: str = "vi") -> Tuple[Optional[BytesIO], str]:
    clean_text = re.sub(r'[<>"\'{}|\\^~\[\]`]', '', text).strip()
    if not clean_text: return None, "Text rỗng."
    chunks = split_text_into_chunks(clean_text); audio_chunks = []
    for chunk in chunks:
        audio = fetch_google_tts_chunk(chunk, lang)
        if audio: audio_chunks.append(audio)
    if not audio_chunks: return None, "Không thể tạo audio."
    return BytesIO(b"".join(audio_chunks)), "ok"

def voice_worker():
    while True:
        try:
            req: VoiceRequest = voice_queue.get(block=True, timeout=1)
            if not req: continue
            voice_text = req.text[:500].strip()
            if not voice_text: voice_queue.task_done(); continue
            audio, result_msg = generate_voice_google(voice_text, req.lang)
            if audio and result_msg == "ok":
                audio.name = f"voice_{int(time.time())}.mp3"
                try: bot.send_voice(req.chat_id, audio, reply_to_message_id=req.reply_id, caption=f"🎙️ {html.escape(voice_text[:200])}", parse_mode="HTML"); brain.stats["voice_generated"] += 1
                except:
                    audio.seek(0)
                    try: bot.send_audio(req.chat_id, audio, reply_to_message_id=req.reply_id, title="Voice", caption=f"🎙️ {html.escape(voice_text[:200])}", parse_mode="HTML"); brain.stats["voice_generated"] += 1
                    except: pass
            else:
                try: bot.send_message(req.chat_id, f"❌ {html.escape(req.user_name)}, không thể tạo giọng nói.\n<i>{result_msg}</i>", reply_to_message_id=req.reply_id, parse_mode="HTML")
                except: pass
            voice_queue.task_done()
        except:
            try: voice_queue.task_done()
            except: pass

for _ in range(4): Thread(target=voice_worker, daemon=True).start()

@bot.message_handler(commands=['voice'])
def voice_cmd(m):
    if not is_grp(m) or antispam(m): return
    users[str(m.from_user.id)] = m.from_user.first_name; save_users(users)
    voice_text = ""
    if m.reply_to_message and m.reply_to_message.text: voice_text = m.reply_to_message.text.strip()
    elif m.text.strip() != '/voice':
        parts = m.text.split(maxsplit=1)
        if len(parts) > 1: voice_text = parts[1].strip()
    if not voice_text: msg = bot.reply_to(m, "❌ /voice [text] hoặc reply.", parse_mode="HTML"); del_msg(m.chat.id, msg.message_id, 10); return
    if len(voice_text) > 500: voice_text = voice_text[:500]
    try:
        voice_queue.put_nowait(VoiceRequest(chat_id=m.chat.id, reply_id=m.message_id, text=voice_text, user_name=m.from_user.first_name))
        msg = bot.reply_to(m, "🎙️ Đang tạo voice...", parse_mode="HTML"); del_msg(m.chat.id, msg.message_id, 5)
    except: bot.reply_to(m, "⚠️ Hàng đợi voice đầy.", parse_mode="HTML")

# ╔══════════════════════════════════════════════════════════════╗
# ║  ĐỌC FILE                                                  ║
# ╚══════════════════════════════════════════════════════════════╝
def detect_encoding(raw_data: bytes) -> str:
    if HAS_CHARDET:
        try:
            result = chardet.detect(raw_data[:50000])
            if result.get('confidence', 0) > 0.5 and result.get('encoding'): return result['encoding']
        except: pass
    for enc in ['utf-8', 'utf-16', 'cp1252', 'latin-1', 'iso-8859-1', 'gbk']:
        try: raw_data.decode(enc); return enc
        except: pass
    return 'utf-8'

def process_file(file_data: bytes, file_name: str, file_id: str) -> Dict:
    if file_id in file_cache and time.time() - file_cache[file_id]["timestamp"] < 300: return file_cache[file_id]
    file_size = len(file_data)
    if file_size > MAX_FILE_SIZE: return {"text": "", "stats": {"error": f"File quá lớn ({file_size/1024/1024:.1f}MB). Tối đa 20MB."}, "summary": "", "preview": "❌ File vượt quá 20MB."}
    
    ext = os.path.splitext(file_name)[1].lower() if file_name else ""; text = ""; stats = {"file_name": file_name, "file_size": file_size, "ext": ext}
    try:
        if ext == '.pdf' and HAS_PYPDF2:
            reader = PyPDF2.PdfReader(BytesIO(file_data))
            text = "\n".join([(p.extract_text() or "") for p in reader.pages]); stats.update({"pages": len(reader.pages), "chars": len(text)})
        elif ext == '.docx' and HAS_DOCX:
            doc = docx.Document(BytesIO(file_data))
            text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()]); stats.update({"paragraphs": len(doc.paragraphs), "chars": len(text)})
        else:
            enc = detect_encoding(file_data); text = file_data.decode(enc, errors='replace'); stats.update({"encoding": enc, "lines": len(text.splitlines()), "chars": len(text)})
    except Exception as e: stats["error"] = str(e)
    
    if len(text) > 100000: preview = text[:3000]; summary = f"File {file_name} ({file_size/1024/1024:.1f}MB) - {len(text):,} ký tự."
    else: preview = text[:5000] if len(text) > 5000 else text; summary = f"File {file_name} ({file_size/1024/1024:.1f}MB) - {len(text):,} ký tự."
    if len(text) > len(preview): preview += f"\n\n... (còn {len(text)-len(preview):,} ký tự)"
    
    result = {"text": text[:100000], "stats": stats, "summary": summary, "preview": preview}
    file_cache[file_id] = {**result, "timestamp": time.time()}
    if len(file_cache) > MAX_CACHE_SIZE:
        oldest = min(file_cache.items(), key=lambda x: x[1]["timestamp"]); del file_cache[oldest[0]]
    brain.stats["files_processed"] += 1
    return result

def search_in_text(text: str, query: str) -> List[str]:
    if not query or not text: return []
    lines = text.splitlines(); results = []
    for i, line in enumerate(lines):
        if query.lower() in line.lower():
            start, end = max(0, i-1), min(len(lines), i+2)
            results.append('\n'.join([f"{j+1}: {lines[j]}" for j in range(start, end)]))
            if len(results) >= 10: break
    return results

@bot.message_handler(content_types=['document'])
def handle_document(m):
    if not is_grp(m) or antispam(m): return
    users[str(m.from_user.id)] = m.from_user.first_name; save_users(users)
    doc = m.document; file_name = doc.file_name or "unknown"; file_size = doc.file_size or 0; file_id = doc.file_id
    if file_size > MAX_FILE_SIZE: msg = bot.reply_to(m, f"❌ File <b>{html.escape(file_name)}</b> quá lớn ({file_size/1024/1024:.1f}MB).", parse_mode="HTML"); del_msg(m.chat.id, msg.message_id, 30); return
    status_msg = bot.reply_to(m, f"📄 Đang đọc <b>{html.escape(file_name)}</b>...", parse_mode="HTML")
    
    def _process():
        try:
            file_info = bot.get_file(file_id); downloaded = bot.download_file(file_info.file_path)
            result = process_file(downloaded, file_name, file_id)
            try: bot.delete_message(m.chat.id, status_msg.message_id)
            except: pass
            if result["stats"].get("error"): bot.reply_to(m, f"❌ Lỗi: {result['stats']['error']}", parse_mode="HTML"); return
            
            stats_text = f"📄 <b>{html.escape(file_name)}</b>\n📏 {file_size/1024/1024:.2f}MB | 📝 {result['stats'].get('chars', 'N/A'):,} ký tự | 📋 {result['stats'].get('lines', 'N/A')} dòng"
            preview = result["preview"]
            if len(preview) > 4000:
                bot.reply_to(m, stats_text, parse_mode="HTML")
                for i in range(0, min(len(preview), 12000), 4000): bot.send_message(m.chat.id, f"<pre>{html.escape(preview[i:i+4000])}</pre>", parse_mode="HTML")
            else:
                bot.reply_to(m, stats_text + f"\n\n<pre>{html.escape(preview)}</pre>", parse_mode="HTML")
        except Exception as e: bot.reply_to(m, f"❌ Lỗi: {str(e)[:200]}", parse_mode="HTML")
    file_executor.submit(_process)

@bot.message_handler(commands=['readfile', 'filesearch'])
def readfile_cmd(m):
    if not is_grp(m): return
    if not m.reply_to_message or not m.reply_to_message.document: msg = bot.reply_to(m, "❌ Reply file + /readfile [từ khóa]", parse_mode="HTML"); del_msg(m.chat.id, msg.message_id, 10); return
    # Gọi lại handle_document với search
    parts = m.text.split(maxsplit=1); query = parts[1].strip() if len(parts) > 1 else ""
    if query:
        doc = m.reply_to_message.document; file_id = doc.file_id
        try:
            file_info = bot.get_file(file_id); downloaded = bot.download_file(file_info.file_path)
            result = process_file(downloaded, doc.file_name or "unknown", file_id)
            search_results = search_in_text(result["text"], query)
            if search_results:
                response = f"🔍 <b>{len(search_results)}</b> kết quả cho '<i>{html.escape(query)}</i>':\n\n"
                for sr in search_results[:5]: response += f"<pre>{html.escape(sr)}</pre>\n---\n"
                if len(search_results) > 5: response += f"\n... và {len(search_results)-5} kết quả khác."
            else: response = f"❌ Không tìm thấy '<i>{html.escape(query)}</i>'."
            bot.reply_to(m, response, parse_mode="HTML")
        except Exception as e: bot.reply_to(m, f"❌ Lỗi: {str(e)[:200]}", parse_mode="HTML")
    else:
        # Gọi document handler
        m.document = m.reply_to_message.document
        m.message_id = m.reply_to_message.message_id
        handle_document(m)

# ╔══════════════════════════════════════════════════════════════╗
# ║  THỐNG KÊ THÀNH VIÊN                                       ║
# ╚══════════════════════════════════════════════════════════════╝
@bot.message_handler(commands=['stats', 'memberstats'])
def stats_cmd(m):
    if not is_grp(m): return
    today = date.today().isoformat(); yesterday = (date.today() - timedelta(days=1)).isoformat()
    last_7_days = [(date.today() - timedelta(days=i)).isoformat() for i in range(7)]
    
    today_join = member_stats["daily_join"].get(today, 0); today_leave = member_stats["daily_leave"].get(today, 0)
    yesterday_join = member_stats["daily_join"].get(yesterday, 0); yesterday_leave = member_stats["daily_leave"].get(yesterday, 0)
    week_join = sum(member_stats["daily_join"].get(d, 0) for d in last_7_days)
    week_leave = sum(member_stats["daily_leave"].get(d, 0) for d in last_7_days)
    
    try: real_count = bot.get_chat_member_count(GROUP_ID)
    except: real_count = member_stats.get("current_members", 0)
    
    text = (f"📊 <b>THỐNG KÊ</b>\n👥 Hiện tại: <b>{real_count}</b> | 📥 Tổng vào: <b>{member_stats['total_joined']}</b> | 📤 Tổng rời: <b>{member_stats['total_left']}</b>\n"
            f"━━━━━━━━━━━━━━━━━━━━\n📅 Hôm nay: ✅ +{today_join} ❌ -{today_leave} 📈 {today_join - today_leave:+d}\n"
            f"📅 Hôm qua: ✅ +{yesterday_join} ❌ -{yesterday_leave} 📈 {yesterday_join - yesterday_leave:+d}\n"
            f"📅 7 ngày: ✅ +{week_join} ❌ -{week_leave} 📈 {week_join - week_leave:+d} (TB: +{week_join/7:.1f}/ngày)")
    msg = bot.reply_to(m, text, parse_mode="HTML"); del_msg(m.chat.id, msg.message_id, 60)

@bot.message_handler(commands=['growth'])
def growth_cmd(m):
    if not is_grp(m): return
    last_7_days = [(date.today() - timedelta(days=i)) for i in range(6, -1, -1)]
    text = "📈 <b>TĂNG TRƯỞNG 7 NGÀY</b>\n<pre>"; max_val = 0
    rows = []
    for d in last_7_days:
        ds = d.isoformat(); join = member_stats["daily_join"].get(ds, 0); leave = member_stats["daily_leave"].get(ds, 0); net = join - leave
        max_val = max(max_val, join, leave, abs(net)); rows.append((d.strftime("%d/%m"), join, leave, net))
    for date_str, join, leave, net in rows:
        join_bar = "█" * int(join/max(max_val,1)*15) if max_val > 0 else ""
        leave_bar = "░" * int(leave/max(max_val,1)*15) if max_val > 0 else ""
        text += f"{date_str} +{join:>3} {join_bar}\n      -{leave:>3} {leave_bar}\n"
    text += "</pre>"; msg = bot.reply_to(m, text, parse_mode="HTML"); del_msg(m.chat.id, msg.message_id, 45)

# ╔══════════════════════════════════════════════════════════════╗
# ║  LỆNH QUẢN LÍ NHÓM                                         ║
# ╚══════════════════════════════════════════════════════════════╝
@bot.message_handler(commands=['ban'])
def ban_cmd(m):
    if not is_grp(m) or not is_admin(m.chat.id, m.from_user.id): return
    target, reason = extract_user_and_reason(m, bot.get_me().username)
    if not target: msg = bot.reply_to(m, "❌ Reply/mention/ID.", parse_mode="HTML"); del_msg(m.chat.id, msg.message_id, 15); return
    try: bot.ban_chat_member(m.chat.id, target); bot.delete_message(m.chat.id, m.message_id); w = bot.send_message(m.chat.id, f"🚫 <b>{html.escape(m.from_user.first_name)}</b> đã ban <code>{target}</code>{' - '+reason if reason else ''}", parse_mode="HTML"); del_msg(m.chat.id, w.message_id, 30)
    except Exception as e: bot.reply_to(m, f"⚠️ Lỗi: {str(e)[:100]}", parse_mode="HTML")

@bot.message_handler(commands=['mute'])
def mute_cmd(m):
    if not is_grp(m) or not is_admin(m.chat.id, m.from_user.id): return
    target, reason = extract_user_and_reason(m, bot.get_me().username)
    if not target: return
    duration = parse_duration(reason) if reason else 3600
    try:
        until = int(time.time()) + duration
        bot.restrict_chat_member(m.chat.id, target, until_date=until, can_send_messages=False, can_send_media_messages=False, can_send_other_messages=False, can_add_web_page_previews=False)
        bot.delete_message(m.chat.id, m.message_id)
        dur_str = f"{duration//3600}h{(duration%3600)//60}m" if duration>=3600 else f"{duration//60}m{duration%60}s"
        w = bot.send_message(m.chat.id, f"🔇 <b>{html.escape(m.from_user.first_name)}</b> mute <code>{target}</code> {dur_str}", parse_mode="HTML"); del_msg(m.chat.id, w.message_id, 30); mutes[target] = until
    except Exception as e: bot.reply_to(m, f"⚠️ {str(e)[:100]}", parse_mode="HTML")

@bot.message_handler(commands=['unmute'])
def unmute_cmd(m):
    if not is_grp(m) or not is_admin(m.chat.id, m.from_user.id): return
    target, _ = extract_user_and_reason(m, bot.get_me().username)
    if not target: return
    try:
        bot.restrict_chat_member(m.chat.id, target, can_send_messages=True, can_send_media_messages=True, can_send_other_messages=True, can_add_web_page_previews=True)
        bot.delete_message(m.chat.id, m.message_id); w = bot.send_message(m.chat.id, f"🔊 Unmute <code>{target}</code>", parse_mode="HTML"); del_msg(m.chat.id, w.message_id, 20)
        if target in mutes: del mutes[target]
    except Exception as e: bot.reply_to(m, f"⚠️ {str(e)[:100]}", parse_mode="HTML")

@bot.message_handler(commands=['warn'])
def warn_cmd(m):
    if not is_grp(m) or not is_admin(m.chat.id, m.from_user.id): return
    target, reason = extract_user_and_reason(m, bot.get_me().username)
    if not target: return
    warn_counts[target] = warn_counts.get(target,0) + 1; cnt = warn_counts[target]
    bot.delete_message(m.chat.id, m.message_id); w = bot.send_message(m.chat.id, f"⚠️ <b>{html.escape(m.from_user.first_name)}</b> warn <code>{target}</code> [{cnt}/3]\n{reason if reason else ''}", parse_mode="HTML"); del_msg(m.chat.id, w.message_id, 25)
    if cnt >= 3:
        try: bot.ban_chat_member(m.chat.id, target, until_date=int(time.time())+3600); del warn_counts[target]
        except: pass

# ╔══════════════════════════════════════════════════════════════╗
# ║  AI + ANTI-SPAM                                            ║
# ╚══════════════════════════════════════════════════════════════╝
def antispam(m) -> bool:
    if is_admin(m.chat.id, m.from_user.id): return False
    uid, now = m.from_user.id, time.time()
    spam[uid] = [t for t in spam.get(uid,[]) if now-t<4] + [now]
    if len(spam[uid]) > 5:
        warn_counts[uid] = warn_counts.get(uid,0) + 1; brain.stats["spam_blocked"] += 1
        try:
            bot.delete_message(m.chat.id, m.message_id)
            if warn_counts[uid] >= 3:
                try: bot.ban_chat_member(m.chat.id, uid, until_date=int(time.time())+3600)
                except: pass
                bot.send_message(m.chat.id, f"🚫 <b>{html.escape(m.from_user.first_name)}</b> bị ban 1h vì spam.", parse_mode="HTML"); del warn_counts[uid]
            else:
                w = bot.send_message(m.chat.id, f"⚠️ Spam {warn_counts[uid]}/3 <b>{html.escape(m.from_user.first_name)}</b>", parse_mode="HTML"); del_msg(m.chat.id, w.message_id, 15)
        except: pass
        return True
    return False

def ask_ai(prompt: str, uid: Optional[int] = None) -> str:
    global ck_idx
    if brain.state == "sleep": return random.choice(get_kho())
    if len(mem) >= 2 and mem[-2] == prompt: return mem[-1]
    sys_msg = "Bạn là kẻ cọc cằn, chửi khịa. Xưng 'tao' gọi 'mày'. Dưới 12 từ."
    msgs = [{"role":"system","content":sys_msg}]
    for txt in list(mem)[-8:]: msgs.append({"role":"user","content":txt})
    msgs.append({"role":"user","content":prompt})
    with ck_lock:
        for _ in range(len(AI_KEYS)):
            k = AI_KEYS[ck_idx]
            if not k["status"] or k["fail"] >= MAX_FAIL: ck_idx = (ck_idx + 1) % len(AI_KEYS); continue
            try:
                resp = ses.post(k["url"], json={"model":k["model"],"messages":msgs,"max_tokens":40,"temperature":0.9}, headers={"Authorization":f"Bearer {k['key']}","Content-Type":"application/json"}, timeout=8)
                if resp.status_code == 200:
                    result = resp.json()['choices'][0]['message']['content'].strip(); result = re.sub(r'[_*`\[\]()]','',result)
                    k["fail"] = 0; k["last_used"] = time.time(); mem.append(prompt); mem.append(result); brain.stats["ai_calls"] += 1; return result
                else: k["fail"] += 1
            except: k["fail"] += 1; brain.stats["errors"] += 1
            ck_idx = (ck_idx + 1) % len(AI_KEYS)
    if not any(k["status"] for k in AI_KEYS):
        for k in AI_KEYS: k["status"], k["fail"] = True, 0
        brain.stats["errors"] = 0; brain.state = "repair"; return "[Não tự sửa] AI đã reset."
    return random.choice(get_kho())

# ╔══════════════════════════════════════════════════════════════╗
# ║  HANDLERS CƠ BẢN                                           ║
# ╚══════════════════════════════════════════════════════════════╝
@bot.message_handler(commands=['start'])
def start(m):
    if not is_grp(m) or antispam(m): return
    users[str(m.from_user.id)] = m.from_user.first_name; save_users(users); brain.trusted_users.add(m.from_user.id)
    help_text = ("<b>🧠 Não Robot - Ultimate</b>\n"
                 "💎 /daily - Điểm danh nhận xu | /balance - Xem xu\n"
                 "🎰 /nohu - Nổ Hũ | /taixiu /baucua /kbb /doanso - Games\n"
                 "📄 /readfile (reply file) - Đọc file 20MB\n"
                 "🎙️ /voice - Text to Speech\n"
                 "📊 /stats /growth /top - Thống kê\n"
                 "🛠️ /ban /mute /unmute /warn - Quản lí\n"
                 "💸 /give [@mention] [xu] - Chuyển xu\n"
                 "Tag/reply bot = chat AI")
    msg = bot.reply_to(m, help_text, parse_mode="HTML"); del_msg(m.chat.id, msg.message_id, 60)

@bot.message_handler(commands=['brain'])
def brain_cmd(m):
    if not is_grp(m): return
    if not is_admin(m.chat.id, m.from_user.id): msg = bot.reply_to(m, "⛔ Không đủ quyền.", parse_mode="HTML"); del_msg(m.chat.id, msg.message_id, 10); return
    uptime = int(time.time() - brain.stats["uptime_start"])
    jackpot = load_jackpot()
    text = (f"🧠 State: <code>{brain.state}</code> | Mood: <code>{brain.mood}</code>\n"
            f"Msgs: <code>{brain.stats['msg_processed']}</code> | Spam: <code>{brain.stats['spam_blocked']}</code>\n"
            f"AI: <code>{brain.stats['ai_calls']}</code> | Voice: <code>{brain.stats['voice_generated']}</code>\n"
            f"Files: <code>{brain.stats['files_processed']}</code> | Daily: <code>{brain.stats['daily_checkins']}</code>\n"
            f"Nổ Hũ spins: <code>{brain.stats['nohu_spins']}</code> | Jackpot: <code>{jackpot:,}</code> xu\n"
            f"Errors: <code>{brain.stats['errors']}</code> | Uptime: <code>{uptime//3600}h{(uptime%3600)//60}m</code>\n"
            f"Voice Q: <code>{voice_queue.qsize()}</code> | Users: <code>{len(users)}</code> | Balances: <code>{len(user_balance)}</code>")
    msg = bot.reply_to(m, text, parse_mode="HTML"); del_msg(m.chat.id, msg.message_id, 30)

@bot.message_handler(func=lambda m: is_grp(m) and m.text)
def handle_text(m):
    if antispam(m) or m.text.startswith('/'): return
    users[str(m.from_user.id)] = m.from_user.first_name; save_users(users)
    brain.think({"uid": m.from_user.id, "txt": m.text, "cmd": False})
    uid = m.from_user.id
    if not brain.should_reply(uid, m.text): return
    if uid in ai_cd and time.time() - ai_cd[uid] < 2: msg = bot.reply_to(m, "Đợi 2s.", parse_mode="HTML"); del_msg(m.chat.id, msg.message_id, 3); return
    ai_cd[uid] = time.time()
    def _ai():
        reply = ask_ai(m.text, uid)
        if f"@{bot.get_me().username}" in m.text or (m.reply_to_message and m.reply_to_message.from_user.id == bot.get_me().id): bot.reply_to(m, html.escape(reply), parse_mode="HTML")
        else: msg_reply = bot.reply_to(m, html.escape(reply), parse_mode="HTML"); del_msg(m.chat.id, msg_reply.message_id)
    ai_executor.submit(_ai)

@bot.message_handler(content_types=['new_chat_members'])
def welcome(m):
    if not is_grp(m): return
    today = date.today().isoformat()
    for u in m.new_chat_members:
        if u.id == bot.get_me().id: continue
        users[str(u.id)] = u.first_name
        if str(u.id) not in member_stats.get("join_dates", {}): member_stats["join_dates"][str(u.id)] = today; member_stats["total_joined"] += 1
        member_stats["daily_join"][today] += 1; member_stats["current_members"] += 1
        save_users(users); save_member_stats()
        msg = bot.send_message(m.chat.id, f"🔥 <a href='tg://user?id={u.id}'>{html.escape(u.first_name)}</a> vừa vào. {random.choice(get_kho())}", parse_mode="HTML"); del_msg(m.chat.id, msg.message_id, 30)

@bot.message_handler(content_types=['left_chat_member'])
def goodbye(m):
    if not is_grp(m): return
    today = date.today().isoformat(); u = m.left_chat_member
    if u.id == bot.get_me().id: return
    member_stats["daily_leave"][today] += 1; member_stats["total_left"] += 1
    member_stats["current_members"] = max(0, member_stats["current_members"] - 1); save_member_stats()
    msg = bot.send_message(m.chat.id, f"🍂 <a href='tg://user?id={u.id}'>{html.escape(u.first_name)}</a> cút. {random.choice(get_kho())}", parse_mode="HTML"); del_msg(m.chat.id, msg.message_id, 30)

# ╔══════════════════════════════════════════════════════════════╗
# ║  BACKGROUND TASKS                                          ║
# ╚══════════════════════════════════════════════════════════════╝
def scheduler_task():
    last_hour = -1; last_midnight = date.today()
    while True:
        try:
            now = datetime.now(tz); brain.health_check()
            if brain.state == "repair": brain.state = "normal"; brain.repair_mode = False
            
            # Reset daily checkin vào 0h
            today = date.today()
            if today != last_midnight:
                last_midnight = today
                # Không reset daily_checkin dict vì mỗi user có ngày riêng
                logger.info(f"Ngày mới: {today}")
            
            if now.minute == 0 and now.hour != last_hour and users:
                uid, uname = random.choice(list(users.items()))
                msg = bot.send_message(GROUP_ID, f"🔔 <b>{now.strftime('%H:%M')}</b> | <a href='tg://user?id={uid}'>{html.escape(uname)}</a>... {random.choice(get_kho())}", parse_mode="HTML"); del_msg(GROUP_ID, msg.message_id, 15); last_hour = now.hour
            if now.minute != 0: last_hour = -1
            
            to_remove = []
            for uid, until in mutes.items():
                if time.time() > until:
                    try: bot.restrict_chat_member(GROUP_ID, uid, can_send_messages=True, can_send_media_messages=True, can_send_other_messages=True, can_add_web_page_previews=True); to_remove.append(uid)
                    except: pass
            for uid in to_remove: del mutes[uid]
            if len(spam) > 100:
                oldest = sorted(spam.items(), key=lambda x: x[1][-1] if x[1] else 0)[:10]
                for uid, _ in oldest: del spam[uid]
        except: pass
        time.sleep(15)

def auto_save_task():
    while True:
        time.sleep(600)
        try: save_users(users); brain.save_state(); save_member_stats(); save_balances(user_balance); save_daily_checkins(daily_checkin); save_jackpot(load_jackpot())
        except: pass

# ╔══════════════════════════════════════════════════════════════╗
# ║  MAIN                                                      ║
# ╚══════════════════════════════════════════════════════════════╝
def main():
    global user_balance, daily_checkin, nohu_jackpot, member_stats
    
    loaded_users = load_users()
    if isinstance(loaded_users, dict): users.update(loaded_users)
    
    user_balance = load_balances()
    daily_checkin = load_daily_checkins()
    nohu_jackpot = load_jackpot()
    
    loaded_stats = load_member_stats()
    member_stats.update(loaded_stats)
    if not isinstance(member_stats.get("daily_join"), defaultdict): member_stats["daily_join"] = defaultdict(int, member_stats.get("daily_join", {}))
    if not isinstance(member_stats.get("daily_leave"), defaultdict): member_stats["daily_leave"] = defaultdict(int, member_stats.get("daily_leave", {}))
    
    try: member_stats["current_members"] = bot.get_chat_member_count(GROUP_ID)
    except: member_stats["current_members"] = len(users)
    
    logger.info(f"🚀 Khởi động: {len(users)} users, {len(user_balance)} balances, Jackpot: {nohu_jackpot:,} xu")
    logger.info(f"📄 PDF={HAS_PYPDF2} DOCX={HAS_DOCX} BS4={HAS_BS4} Chardet={HAS_CHARDET}")
    logger.info(f"🎰 Nổ Hũ | 💎 Điểm Danh | 🎲 Games | 🎙️ Voice | 📊 Stats | 🛠️ Admin")
    
    Thread(target=scheduler_task, daemon=True).start()
    Thread(target=auto_save_task, daemon=True).start()
    
    try: bot.infinity_polling(timeout=30, none_stop=True, interval=0.5)
    except Exception as e:
        logger.critical(f"Bot dừng: {e}")
        brain.stats["errors"] += 1; brain.save_state(); save_balances(user_balance); save_daily_checkins(daily_checkin); save_jackpot(load_jackpot()); save_member_stats()

if __name__ == "__main__":
    main()
